import openpyxl
import docx
from openpyxl.utils import get_column_letter


# Додатковий клас речей в чеку 
class Item:
    def __init__(self, id1):
        self.item_id = id1
        self.name = ''
        self.units = ''
        self.price = 0
        self.quantity = 0
        
# Головний клас в якому зберігається вся інформація про чек 
class Check:
    def __init__(self, id1):
        self.check_id = id1
        self.number = 0
        self.date1 = ''
        self.client_id = 0
        self.client_name = ''
        self.client_address = ''
        self.items_in_check = [] # id of item, names, units, price and quantity
    
# вікриваємо файл
wb = openpyxl.load_workbook("info.xlsx")
# масив чеків
arr_of_checks = []

# Проходимося по кожному чеку  та закидуємо всю інформацію яка зв'язана з чеком з інших аркушів
for row in range(2, 1000):
    
    wchecks = wb["checks"]
    col = get_column_letter(1)
    if(wchecks[col + str(row)].value == None):
        break
    id1 = wchecks[col + str(row)].value
    check = Check(id1)
    print("check id is: " + check.check_id)
    col = get_column_letter(2)
    check.number = wchecks[col + str(row)].value
    print("number is: " + str(check.number))
    col = get_column_letter(3)
    check.date1 = wchecks[col + str(row)].value
    print("date is: " + str(check.date1))
    col = get_column_letter(4)
    check.client_id = wchecks[col + str(row)].value
    print("client id is: " + check.client_id)
    
    wbuyers = wb["buyers"]
    for buyers_row in range(2, 1000):    
        col = get_column_letter(1)
        # print(wbuyers[col + str(buyers_row)].value)
        if(wbuyers[col + str(buyers_row)].value == None):
            break
        elif(wbuyers[col + str(buyers_row)].value == check.client_id):
            
            col = get_column_letter(2)
            check.client_name = wbuyers[col + str(buyers_row)].value
            print("client name is: " + check.client_name)
            col = get_column_letter(3)
            check.client_address = wbuyers[col + str(buyers_row)].value
            print("client address is: " + check.client_address)
            break
        
    print()
    witems = wb["items_in_checks"]
    for items_row in range(2, 1000):

        col = get_column_letter(1)
        if(witems[col + str(items_row)].value == None):
            break
        elif(witems[col + str(items_row)].value == check.check_id):
            
            col = get_column_letter(2)
            item = Item(witems[col + str(items_row)].value)
            print("item id is: " + item.item_id)
            col = get_column_letter(3)
            item.quantity = witems[col + str(items_row)].value
            print("quantity of item is: " + str(item.quantity))
            
            wproducts = wb["products"]
            for products_row in range(2, 1000):
                col = get_column_letter(1)
                if(wproducts[col + str(products_row)].value == None):
                    break
                elif(wproducts[col + str(products_row)].value == item.item_id):
                    
                    col = get_column_letter(2)
                    item.name = wproducts[col + str(products_row)].value
                    print("name of item is: " + item.name)
                    col = get_column_letter(3)
                    item.units = wproducts[col + str(products_row)].value
                    print("units of item is: " + item.units)
                    col = get_column_letter(4)
                    item.price = wproducts[col + str(products_row)].value
                    print("item price is: " + str(item.price))
                    check.items_in_check.append(item)
                    break
                
            print()
    arr_of_checks.append(check)
    print()


# Створюємо документ для кожного чеку в який закидуємо інформацію яку просить умова
information = ["№", "Назва", "Од.виміру", "Кількість", "Ціна","Сума"]
for check in arr_of_checks:
    doc = docx.Document()
        
    doc.add_paragraph("Рахунок №" + str(int(check.number)) + " " * 100 + "Дата " + str(check.date1).split(" ")[0])
    doc.add_paragraph("Покупець: " + check.client_name)
    
    total = 0
    table = doc.add_table(rows=1, cols=6)
    table.style = "TableGrid"
    for i in range(6):
            table.rows[0].cells[i].text = information[i]
    i = 1
    for item in check.items_in_check:
        total += item.price * item.quantity
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(item.name)
        cells[2].text = str(item.units)
        cells[3].text = str(int(item.quantity))
        cells[4].text = str(float(item.price))
        cells[5].text = str(float(item.price * item.quantity))
        i += 1
        
    doc.add_paragraph(" " * 130 + "Всього: " + str(float(total)))
    doc.save("check_" + str(check.check_id) + ".docx")